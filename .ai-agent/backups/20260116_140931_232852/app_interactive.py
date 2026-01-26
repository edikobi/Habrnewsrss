from app.database import (
    get_db_session, User, ContentItem, UserInterest, 
)
from app.database import UserSettings, FavoriteContent
from datetime import datetime

import os
import re
import datetime
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
from sqlalchemy import func
import click
from colorama import init, Fore, Back, Style
from docx import Document

from app.database import (
    get_db_session, User, ContentItem, UserInterest, 
    UserProgress, Tag, SessionLocal, UserSettings, FavoriteContent
)
from app.services.aggregator import ContentAggregator, update_all_content
from app.services.recommender import ContentRecommender
from app.services.email_sender import EmailSender
from app.config import settings
from app.sources.youtube import YouTubeSource
from app.sources.habr import HabrSource
from app.sources.coursera import CourseraSource

logger = logging.getLogger(__name__)


class InteractiveMenu:
    """Интерактивное меню для системы агрегации контента с управлением источниками и поиском."""

    def __init__(self):
        self.session = SessionLocal()
        self.aggregator = ContentAggregator(db_session=self.session)
        self.recommender = ContentRecommender(db_session=self.session)
        self.email_sender = EmailSender()
        self.sources = {
            "youtube": {"enabled": bool(settings.youtube_api_key), "instance": YouTubeSource()},
            "habr": {"enabled": True, "instance": HabrSource()},
            "coursera": {"enabled": bool(settings.coursera_api_key), "instance": CourseraSource()},
        }
        self.current_user_id: Optional[int] = None

        # Load saved user ID if available
        try:
            saved_id = self._load_current_user_id()
            if saved_id is not None:
                self.current_user_id = saved_id
        except Exception as e:
            logger.warning(f"Failed to load saved user ID: {e}")

    def _get_current_user_file_path(self) -> Any:
        """Get path to .current_user file in project root directory."""
        return Path.cwd() / '.current_user'

    def _load_current_user_id(self) -> Optional[int]:
        """Load current user ID from .current_user file if it exists and contains valid integer."""
        try:
            path = self._get_current_user_file_path()
            if path.exists():
                content = path.read_text(encoding='utf-8').strip()
                if content and content.isdigit():
                    return int(content)
            return None
        except Exception as e:
            logger.warning(f"Failed to load current user ID: {e}")
            return None

    def _save_current_user_id(self, user_id: Optional[int]) -> None:
        """Save current user ID to .current_user file. If user_id is None, delete the file."""
        try:
            path = self._get_current_user_file_path()
            if user_id is None or user_id <= 0:
                path.unlink(missing_ok=True)
                return
            path.write_text(str(user_id), encoding='utf-8')
        except Exception as e:
            logger.warning(f"Failed to save current user ID: {e}")

    def _ensure_demo_data(self) -> None:
        """Обеспечить наличие демо-пользователя и контента."""
        try:
            user_count = self.session.query(User).count()
            if user_count == 0:
                user = User(username="demo", email="demo@example.com")
                self.session.add(user)
                self.session.flush()

                # Create default settings
                settings_obj = UserSettings(user_id=user.id, email_digest=user.email)
                self.session.add(settings_obj)

                for tag_name in ["python", "programming", "artificial intelligence"]:
                    interest = UserInterest(user_id=user.id, tag_name=tag_name)
                    self.session.add(interest)

                self.session.commit()
                print(Fore.GREEN + f"✓ Создан демо-пользователь с ID: {user.id}")

            content_count = self.session.query(ContentItem).count()
            if content_count == 0 and self.sources["habr"]["enabled"]:
                items = self.sources["habr"]["instance"].fetch_content(["python"], 5)
                self.aggregator.save_content_items(items)
                print(Fore.GREEN + f"✓ Добавлено {len(items)} демо-элементов контента из Habr")
        except Exception as e:
            print(Fore.YELLOW + f"⚠ Создание демо-данных пропущено: {e}")

    def add_new_user_interactive(self) -> None:
        """Interactively add a new user with username and email, creating default settings."""
        print(Fore.CYAN + "\n" + "═"*40)
        print(Fore.CYAN + "ДОБАВЛЕНИЕ НОВОГО ПОЛЬЗОВАТЕЛЯ")
        print(Fore.CYAN + "─"*40)

        username = input(Fore.YELLOW + "Введите имя пользователя: ").strip()
        if not username:
            print(Fore.RED + "Имя пользователя не может быть пустым.")
            return

        email = input(Fore.YELLOW + "Введите email: ").strip()
        if not email or "@" not in email:
            print(Fore.RED + "Некорректный email.")
            return

        try:
            existing = self.session.query(User).filter(
                (User.username == username) | (User.email == email)
            ).first()

            if existing:
                print(Fore.RED + "Пользователь с таким именем или email уже существует.")
                return

            user = User(username=username, email=email)
            self.session.add(user)
            self.session.flush()

            settings_obj = UserSettings(user_id=user.id, email_digest=email)
            self.session.add(settings_obj)

            self.session.commit()
            self.current_user_id = user.id

            # Save selection to file
            try:
                self._save_current_user_id(user.id)
            except Exception as e:
                logger.warning(f"Failed to save user selection: {e}")
            print(Fore.GREEN + f"✓ Пользователь создан! ID: {user.id}")
            print(Fore.YELLOW + "Теперь вы можете настроить интересы в меню настроек.")
        except Exception as e:
            self.session.rollback()
            print(Fore.RED + f"Ошибка при создании пользователя: {e}")

    def run(self) -> None:
        """Основной интерактивный цикл."""
        init(autoreset=True)
        self._ensure_demo_data()

        # Проверка пропущенных обновлений контента при запуске
        try:
            click.echo(Fore.CYAN + "Проверка пропущенных обновлений контента...")
            aggregator = ContentAggregator()
            results = aggregator.check_and_update_content_for_all_users()
            if results:
                total = sum(results.values())
                click.echo(Fore.GREEN + f"✓ Добавлено {total} пропущенных статей в базу данных")
                for user_id, count in results.items():
                    if count > 0:
                        click.echo(Fore.GREEN + f"  Пользователь {user_id}: {count} элементов")
            else:
                click.echo(Fore.YELLOW + "✓ Нет пропущенных обновлений контента")
        except Exception as e:
            logger.error(f"Ошибка проверки пропущенных обновлений контента: {e}")
            click.echo(Fore.RED + f"⚠ Ошибка проверки обновлений: {e}")

        # Check for missed digests on startup
        try:
            missed = self.aggregator.check_missed_digests()
            if any(missed.values()):
                print(Fore.GREEN + "✓ Отправлены пропущенные подборки на email")
        except Exception as e:
            logger.error(f"Error checking missed digests: {e}")

        # Try to restore last user from saved file
        try:
            if self.current_user_id is None:
                users = self.session.query(User).all()
                if users and len(users) == 1:
                    self.current_user_id = users[0].id
                    print(Fore.YELLOW + f"Автоматически выбран пользователь: {users[0].username}")
        except Exception as e:
            logger.error(f"Error restoring user: {e}")

        print(Fore.CYAN + "\n" + "="*50)
        print(Fore.CYAN + "АГРЕГАТОР ОБРАЗОВАТЕЛЬНОГО КОНТЕНТА - ИНТЕРАКТИВНЫЙ РЕЖИМ")
        print(Fore.CYAN + "="*50)

        try:
            while True:
                choice = self.show_main_menu()
                if choice == 8:
                    print(Fore.CYAN + "\nДо свидания!")
                    break
        except KeyboardInterrupt:
            print(Fore.CYAN + "\n\nДо свидания!")
        finally:
            self.session.close()

    def show_main_menu(self) -> int:
        """Показать главное меню и обработать выбор пользователя."""
        print(Fore.CYAN + "\n" + "═"*40)
        print(Fore.CYAN + "ГЛАВНОЕ МЕНЮ")
        print(Fore.CYAN + "─"*40)

        if self.current_user_id:
            user = self.session.query(User).get(self.current_user_id)
            if user:
                print(f"{Fore.GREEN}Текущий пользователь: {user.username} (ID: {user.id})")

        print(f"{Fore.YELLOW}1.{Fore.WHITE} Просмотр контента и новостей")
        print(f"{Fore.YELLOW}2.{Fore.WHITE} Управление источниками")
        print(f"{Fore.YELLOW}3.{Fore.WHITE} Поиск контента")
        print(f"{Fore.YELLOW}4.{Fore.WHITE} Получить рекомендации")
        print(f"{Fore.YELLOW}5.{Fore.WHITE} Настройки пользователя")
        print(f"{Fore.YELLOW}6.{Fore.WHITE} Просмотр статистики")
        print(f"{Fore.YELLOW}7.{Fore.WHITE} Просмотр всех статей из БД")
        print(f"{Fore.YELLOW}8.{Fore.WHITE} Выход")
        print(Fore.CYAN + "─"*40)

        while True:
            try:
                choice = int(input(Fore.YELLOW + "\nВведите ваш выбор (1-8): "))
                if 1 <= choice <= 8:
                    break
                print(Fore.RED + "Пожалуйста, введите число от 1 до 8.")
            except ValueError:
                print(Fore.RED + "Пожалуйста, введите корректное число.")

        if choice == 1:
            self.view_content()
        elif choice == 2:
            self.manage_sources()
        elif choice == 3:
            self.search_content()
        elif choice == 4:
            self.get_recommendations_interactive()
        elif choice == 5:
            self.configure_user_settings()
        elif choice == 6:
            self.show_statistics()
        elif choice == 7:
            self.view_all_content_from_db()

        return choice

    def view_content(self) -> None:
        """Просмотр и поиск контента с возможностью добавления в избранное."""
        print(Fore.CYAN + "\n" + "═"*40)
        print(Fore.CYAN + "ПРОСМОТР КОНТЕНТА")
        print(Fore.CYAN + "─"*40)

        if not self.current_user_id:
            users = self.session.query(User).all()
            if users:
                print(Fore.CYAN + "\nВыберите пользователя:")
                for u in users:
                    print(f"ID: {u.id} | {u.username}")
                uid = input(Fore.YELLOW + "Введите ID или нажмите Enter для пропуска: ")
                if uid.isdigit():
                    self.current_user_id = int(uid)
                else:
                    print(Fore.YELLOW + "Поиск без пользователя...")

        user_id = self.current_user_id

        if user_id:
            try:
                user = self.session.query(User).get(user_id)
                if not user:
                    print(Fore.RED + f"Пользователь с ID {user_id} не найден.")
                    return

                print(Fore.WHITE + f"\nГенерация ежедневной подборки для пользователя: {Fore.GREEN}{user.username}")
                digest = self.aggregator.get_daily_digest(user_id, max_items=15)

                if not digest:
                    print(Fore.YELLOW + "Контент не найден.")
                    return

                print(Fore.GREEN + f"\nНайдено {len(digest)} элементов:")
                for i, item in enumerate(digest, 1):
                    pub_date = item.published_at.strftime('%Y-%m-%d %H:%M') if item.published_at else "Неизвестно"
                    add_date = item.added_at.strftime('%Y-%m-%d %H:%M') if item.added_at else "Неизвестно"
                    print(f"{Fore.YELLOW}{i}.{Fore.WHITE} {item.title}")
                    print(f"   Платформа: {Fore.CYAN}{item.platform}{Fore.WHITE} | Сложность: {Fore.CYAN}{item.difficulty}")
                    print(f"   Опубликовано: {pub_date} | Добавлено: {add_date}")
                    print(f"   URL: {Fore.BLUE}{item.url}")
                    print()

                choice = input(Fore.YELLOW + "Введите номер для деталей/отметки или 0 для выхода: ")
                if choice.isdigit() and 0 < int(choice) <= len(digest):
                    self.view_content_details(digest[int(choice)-1].id)

            except Exception as e:
                print(Fore.RED + f"Ошибка: {e}")
        else:
            print(Fore.WHITE + "\nНедавний контент из всех источников:")
            items = self.session.query(ContentItem).order_by(ContentItem.published_at.desc()).limit(20).all()

            for i, item in enumerate(items, 1):
                pub_date = item.published_at.strftime('%Y-%m-%d %H:%M') if item.published_at else "Неизвестно"
                add_date = item.added_at.strftime('%Y-%m-%d %H:%M') if item.added_at else "Неизвестно"
                print(f"{Fore.YELLOW}{i}.{Fore.WHITE} {item.title}")
                print(f"   Платформа: {Fore.CYAN}{item.platform}{Fore.WHITE} | Опубликовано: {pub_date}")
                print(f"   Добавлено: {add_date}")
                print(f"   URL: {Fore.BLUE}{item.url}")
                print()

    def view_content_details(self, content_id: int) -> None:
        """Display detailed information about a specific content item."""
        try:
            item = self.session.query(ContentItem).get(content_id)
            if not item:
                print(Fore.RED + "Контент не найден.")
                return

            print(Fore.CYAN + "\n" + "═"*40)
            print(Fore.WHITE + f"ЗАГОЛОВОК: {Fore.GREEN}{item.title}")
            print(Fore.WHITE + f"ПЛАТФОРМА: {Fore.CYAN}{item.platform}")
            print(Fore.WHITE + f"URL: {Fore.BLUE}{item.url}")
            print(Fore.WHITE + f"СЛОЖНОСТЬ: {Fore.CYAN}{item.difficulty}")

            pub_date = item.published_at.strftime('%Y-%m-%d %H:%M') if item.published_at else "Неизвестно"
            add_date = item.added_at.strftime('%Y-%m-%d %H:%M') if item.added_at else "Неизвестно"
            print(Fore.WHITE + f"ОПУБЛИКОВАНО: {pub_date}")
            print(Fore.WHITE + f"ДОБАВЛЕНО: {add_date}")

            if item.full_text and item.full_text.strip():
                print(Fore.WHITE + "\nПОЛНЫЙ ТЕКСТ СТАТЬИ:")
                full_text_len = len(item.full_text)
                chunk_size = 3500
                total_pages = (full_text_len + chunk_size - 1) // chunk_size
                current_page = 0

                while True:
                    start = current_page * chunk_size
                    end = min(start + chunk_size, full_text_len)
                    print(Fore.WHITE + item.full_text[start:end])
                    print(Fore.CYAN + f"\n[Страница {current_page + 1} из {total_pages}]")
                    print(Fore.YELLOW + "n - следующая | p - предыдущая | число - перейти | q - выход")
                    page_choice = input(Fore.YELLOW + "Навигация: ").lower().strip()
                    if page_choice == 'n' and current_page < total_pages - 1:
                        current_page += 1
                    elif page_choice == 'p' and current_page > 0:
                        current_page -= 1
                    elif page_choice == 'q':
                        break
                    elif page_choice.isdigit():
                        p_num = int(page_choice)
                        if 1 <= p_num <= total_pages:
                            current_page = p_num - 1
            else:
                print(Fore.WHITE + "\nОПИСАНИЕ:")
                print(Fore.WHITE + (item.description or "Нет описания"))

            tags = ", ".join([t.name for t in item.tags])
            print(Fore.WHITE + f"\nТЕГИ: {Fore.YELLOW}{tags}")
            print(Fore.CYAN + "─"*40)

            if self.current_user_id:
                print(f"{Fore.YELLOW}1.{Fore.WHITE} Отметить как завершенное")
                print(f"{Fore.YELLOW}2.{Fore.WHITE} Добавить в избранное")
                print(f"{Fore.YELLOW}3.{Fore.WHITE} Экспортировать статью в MD")
                print(f"{Fore.YELLOW}4.{Fore.WHITE} Экспортировать статью в DOCX")
                print(f"{Fore.YELLOW}5.{Fore.WHITE} Назад к списку")

                choice = input(Fore.YELLOW + "\nВыбор: ")
                if choice == "1":
                    rating = input(Fore.YELLOW + "Оценка (1-5, по умолчанию 5): ")
                    rating = int(rating) if rating.isdigit() else 5
                    notes = input(Fore.YELLOW + "Заметки: ")
                    progress = UserProgress(user_id=self.current_user_id, content_id=item.id)
                    progress.mark_completed(rating=rating, notes=notes)
                    self.session.add(progress)
                    self.session.commit()
                    print(Fore.GREEN + "✓ Отмечено как завершенное!")
                elif choice == "2":
                    fav = FavoriteContent(user_id=self.current_user_id, content_id=item.id)
                    self.session.add(fav)
                    self.session.commit()
                    print(Fore.GREEN + "✓ Добавлено в избранное!")
                elif choice == "3":
                    self.export_article_to_md(item)
                elif choice == "4":
                    self.export_article_to_docx(item)
        except Exception as e:
            print(Fore.RED + f"Ошибка при просмотре деталей: {e}")

    def export_article_to_md(self, item: ContentItem) -> None:
        """Export article content to Markdown file with metadata and full text."""
        try:
            # Create export directory in project root
            # interactive.py -> app -> root
            project_root = Path(__file__).resolve().parent.parent
            export_dir = project_root / "exports"
            export_dir.mkdir(exist_ok=True)

            # Sanitize filename from title
            safe_title = re.sub(r'[^\w\s-]', '', item.title)[:50]
            safe_title = safe_title.replace(' ', '_')
            filename = f"{item.id}_{safe_title}.md"
            filepath = export_dir / filename

            # Build Markdown content
            md = f"# {item.title}\n\n"
            md += "## Метаданные\n\n"
            md += f"- **Платформа:** {item.platform}\n"
            md += f"- **Ссылка:** {item.url}\n"
            md += f"- **Сложность:** {item.difficulty.value if item.difficulty else 'N/A'}\n"

            pub_date = item.published_at.strftime('%Y-%m-%d %H:%M') if item.published_at else 'Неизвестно'
            add_date = item.added_at.strftime('%Y-%m-%d %H:%M') if item.added_at else 'Неизвестно'
            md += f"- **Опубликовано:** {pub_date}\n"
            md += f"- **Добавлено в БД:** {add_date}\n"

            tags = ", ".join([t.name for t in item.tags]) if item.tags else "Нет тегов"
            md += f"- **Теги:** {tags}\n\n"

            md += "## Содержание\n\n"
            if item.full_text and item.full_text.strip():
                md += item.full_text + "\n"
            else:
                md += (item.description or "Нет содержания") + "\n"

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(md)

            print(Fore.GREEN + f"✓ Статья экспортирована: {filepath}")
        except Exception as e:
            print(Fore.RED + f"Ошибка при экспорте: {e}")

    def export_article_to_docx(self, item: ContentItem) -> None:
        """Export article content to DOCX file in the project root exports folder."""
        try:
            # interactive.py -> app -> root
            project_root = Path(__file__).resolve().parent.parent
            export_dir = project_root / "exports"
            export_dir.mkdir(exist_ok=True)

            safe_title = re.sub(r'[^\w\s-]', '', item.title)[:50]
            safe_title = safe_title.replace(' ', '_')
            filename = f"{item.id}_{safe_title}.docx"
            filepath = export_dir / filename

            doc = Document()
            doc.add_heading(item.title, level=1)

            doc.add_heading("Метаданные", level=2)
            doc.add_paragraph(f"Платформа: {item.platform}")
            doc.add_paragraph(f"URL: {item.url}")
            doc.add_paragraph(f"Сложность: {item.difficulty.value if item.difficulty else 'N/A'}")

            pub_date = item.published_at.strftime('%Y-%m-%d %H:%M') if item.published_at else 'Неизвестно'
            add_date = item.added_at.strftime('%Y-%m-%d %H:%M') if item.added_at else 'Неизвестно'
            doc.add_paragraph(f"Опубликовано: {pub_date}")
            doc.add_paragraph(f"Добавлено в БД: {add_date}")

            tags = ", ".join([t.name for t in item.tags]) if item.tags else "Нет тегов"
            doc.add_paragraph(f"Теги: {tags}")

            doc.add_heading("Содержание", level=2)
            content = item.full_text if item.full_text and item.full_text.strip() else (item.description or "Нет содержания")
            doc.add_paragraph(content)

            doc.save(str(filepath))
            print(Fore.GREEN + f"✓ Статья экспортирована в DOCX: {filepath}")

        except Exception as e:
            print(Fore.RED + f"Ошибка при экспорте в DOCX: {e}")

    def manage_sources(self) -> None:
        """Включить/выключить и настроить источники."""
        print(Fore.CYAN + "\n" + "═"*40)
        print(Fore.CYAN + "УПРАВЛЕНИЕ ИСТОЧНИКАМИ")
        print(Fore.CYAN + "─"*40)

        while True:
            print(Fore.WHITE + "\nТекущий статус источников:")
            for source_name, source_info in self.sources.items():
                status = f"{Fore.GREEN}ВКЛЮЧЕН" if source_info["enabled"] else f"{Fore.RED}ВЫКЛЮЧЕН"
                print(f"  {Fore.WHITE}{source_name.upper():10} : {status}")

            print(Fore.WHITE + "\nОпции:")
            print(f"{Fore.YELLOW}1.{Fore.WHITE} Включить/выключить источник")
            print(f"{Fore.YELLOW}2.{Fore.WHITE} Настроить API ключи")
            print(f"{Fore.YELLOW}3.{Fore.WHITE} Установить лимиты контента")
            print(f"{Fore.YELLOW}4.{Fore.WHITE} Проверить подключение источника")
            print(f"{Fore.YELLOW}5.{Fore.WHITE} Вернуться в главное меню")

            try:
                choice = int(input(Fore.YELLOW + "\nВведите выбор (1-5): "))
                if choice == 5:
                    break
                if choice == 1:
                    source_name = input(Fore.YELLOW + "Введите название источника: ").lower()
                    if source_name in self.sources:
                        self.sources[source_name]["enabled"] = not self.sources[source_name]["enabled"]
                elif choice == 4:
                    source_name = input(Fore.YELLOW + "Введите название источника для теста: ").lower()
                    if source_name in self.sources and self.sources[source_name]["enabled"]:
                        source = self.sources[source_name]["instance"]
                        test_items = source.fetch_content(["python"], 1)
                        print(Fore.GREEN + f"✓ Найдено {len(test_items)} элементов.")
            except Exception as e:
                print(Fore.RED + f"Ошибка: {e}")

    def search_content(self) -> None:
        """Интерактивный поиск."""
        keywords_input = input(Fore.YELLOW + "Введите ключевые слова через запятую: ")
        keywords = [k.strip() for k in keywords_input.split(",") if k.strip()]
        if not keywords: return
        
        results = self.aggregator.search_content(keywords)
        for i, item in enumerate(results, 1):
            print(f"{Fore.YELLOW}{i}.{Fore.WHITE} {item.title} ({item.platform})")

    def get_recommendations_interactive(self) -> None:
        """Интерактивное меню для получения рекомендаций."""
        if not self.current_user_id: return
        recs = self.recommender.get_recommendations(self.current_user_id, 5)
        for i, item in enumerate(recs, 1):
            print(f"{Fore.YELLOW}{i}.{Fore.WHITE} {item.title} ({item.platform})")

    def configure_user_settings(self) -> None:
        """Меню настроек пользователя."""
        print(f"{Fore.YELLOW}1.{Fore.WHITE} Выбор пользователя")
        print(f"{Fore.YELLOW}2.{Fore.WHITE} Настройки email")
        print(f"{Fore.YELLOW}3.{Fore.WHITE} Управление интересами")
        print(f"{Fore.YELLOW}4.{Fore.WHITE} Назад")
        choice = input(Fore.YELLOW + "\nВыбор: ")
        if choice == "1":
            users = self.session.query(User).all()
            for u in users: print(f"ID: {u.id} | {u.username}")
            uid = input("Введите ID: ")
            if uid.isdigit():
                self.current_user_id = int(uid)
                self._save_current_user_id(self.current_user_id)

    def show_statistics(self) -> None:
        """Показать статистику."""
        total_items = self.session.query(ContentItem).count()
        print(f"Всего элементов: {total_items}")

    def view_all_content_from_db(self) -> None:
        """Просмотр всех статей из БД."""
        items = self.session.query(ContentItem).limit(20).all()
        for i, item in enumerate(items, 1):
            print(f"{Fore.YELLOW}{i}.{Fore.WHITE} {item.title} [{item.platform}]")